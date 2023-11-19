import os
import sys
import random
from docx import Document
import datetime
import uuid

from PIL import Image as PILImage, ImageDraw
import random

import pandas as pd
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.chart import BarChart,Reference

import numpy as np

#Reportlab to prepare the PDFs
from reportlab.platypus import SimpleDocTemplate
from reportlab.platypus import Paragraph, Spacer, Table, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

#fake intellisense
from commonSeedDataTypes import commonSeedDataTypes as cSDY
from commonSeedDataTypes import commonCityDictValues as cCityFields

# our code
import libFakeData as fd

#make a semi unique batch name so there shouldn't be collisions 
#if we make multiple uploads of gibberish data.
batchName = str(uuid.uuid4())[:8]

def loadSeedData():
    global sdFirstNames
    global sdLastNames 
    global sdLastNames 
    global sdStreetNames 
    global sdCourseNames 
    global sdCompanyNames 
    global sdTextParts 
    global sdApplicationTypes
    global sdPurchaseObjects 
    global sdContractServices 
    global sdCityDictList 
    global sdCreditCards 
    global sdListOfDogPictures

    sdFirstNames = fd.getSeedDataFromFile(cSDY.FirstNames)
    sdLastNames = fd.getSeedDataFromFile(cSDY.LastNames)
    sdLastNames = fd.getSeedDataFromFile(cSDY.LastNames)
    sdStreetNames = fd.getSeedDataFromFile(cSDY.StreetNames)
    sdCourseNames = fd.getSeedDataFromFile(cSDY.CourseNames)
    sdCompanyNames = fd.getSeedDataFromFile(cSDY.CompanyNames)
    sdTextParts = fd.getSeedDataFromFile(cSDY.TextParts)
    sdApplicationTypes = fd.getSeedDataFromFile(cSDY.ApplicationTypes)
    sdPurchaseObjects = fd.getSeedDataFromFile(cSDY.PurchaseObjects)
    sdContractServices = fd.getSeedDataFromFile(cSDY.ContractServices)
    sdCreditCards = fd.getSeedDataFromFile(cSDY.CreditCards)
    sdCityDictList = fd.getSeedDataCityList()
    sdListOfDogPictures = fd.getListOfDogPictures()

currentPath = os.path.dirname(os.path.abspath(__file__))
outPutPath = os.path.join(currentPath,"OutPut") 

if not os.path.exists(outPutPath):
    print("Output path does NOT exists ", outPutPath)
    os.mkdir(outPutPath)
else:
    print("Output path exists ", outPutPath)

def makeRandomMoodFace():
    colours = ["yellow","green","red","blue","brown","pink", "purple","teal","grey"]
    moods = ["happy","sad","neutral","suprized"]

    # Create a new image with a white background
    img = PILImage.new("RGB", (500, 500), "white")

    # Create a draw object to draw on the image
    draw = ImageDraw.Draw(img)

    # Draw a random coloured circle for the face
    draw.ellipse((100, 100, 400, 400), fill=random.choice(colours))

    # Draw two black circles for the eyes
    draw.ellipse((175, 175, 225, 225), fill="black")
    draw.ellipse((275, 175, 325, 225), fill="black")

    # Draw a black arc for the mood
    yStart = random.choice(range(250,300))

    mood = random.choice(moods)
    if mood == "happy":
        draw.arc((175, yStart, 325, yStart+100), start=340, end=200, fill="black", width=10)
    elif mood == "suprized":
        draw.arc((175, yStart, 325, yStart+100), start=0, end=360, fill="black", width=10)
    elif mood == "sad":
        draw.arc((175, yStart, 325, yStart+100), start=200, end=340, fill="black", width=10)
    else:
        # must be neutral
        draw.line((175, yStart, 325, yStart+100),fill="black",width=10)

    # Save the image
    img.save("./SeedData/moodface.png")

def makeBuildingSitePhoto(documentName=""):
    colours = ["yellow","green","red","blue","brown","pink", "purple","teal","grey"]
 
    # Create a new image with a white background
    img = PILImage.new("RGB", (500, 500), "white")

    boxHeight = random.choice(range(200,400))
    boxWidth = random.choice(range(100,300))

    # Create a draw object to draw on the image
    draw = ImageDraw.Draw(img)

    # Draw a random coloured circle for the face
    #draw.rectangle(xy=((1,1,500,500)), fill="black")
    draw.rectangle(xy=((10,boxHeight),(10+boxWidth,500)), fill=random.choice(colours))

    #door
    draw.rectangle(xy=((20,450),(40,500)), fill="white")

    draw.polygon(xy=((0,boxHeight),(boxWidth+20,boxHeight), ((boxWidth)/2,boxHeight-100)), fill = random.choice(colours))

    # Save the image
    img.save(f"./OutPut/{documentName}.jpeg")

def outputContracts(numberToPrint = 1):
    progressCount = 0
    for sequence in range (0,numberToPrint):
        date = datetime.datetime.now()+datetime.timedelta(weeks= -1*random.choice(range(1,300,1)))
        year =  date.strftime("%Y")
        month = date.strftime("%b")
        batch = date.strftime("")
        documentName = f"C{year}-{month}.{sequence}B{batchName}.docx"
        #Contract 
        msDoc = Document()
        msDoc.add_heading(f"Contract {year}-{month}.{sequence}")
        company = fd.getRandomItemFromList(sdCompanyNames)
        paragraph = msDoc.add_paragraph(f"This is a contract between us and {company} for {fd.getRandomItemFromList(sdTextParts)}.")
        tableRowCount = 5
        table = msDoc.add_table(rows=tableRowCount, cols=4)
        #header row
        row_cells = table.rows[0].cells
        row_cells[0].text = "item"
        row_cells[1].text = "quantity"
        row_cells[2].text = "price"
        row_cells[3].text = "" 
        total = 0.00

        for i in range(1,tableRowCount-1):
            #add row
            row_cells = table.rows[i].cells
            row_cells[0].text = fd.getRandomItemFromList(sdPurchaseObjects)
            quantity = random.choice([2,10,12,42])
            row_cells[1].text = str(quantity)
            unitPrice = random.choice([1.00,3.00,7.00,4.20])
            row_cells[2].text = str(unitPrice) 
            rowPrice = unitPrice*quantity
            row_cells[3].text = str("{:.2f}".format(unitPrice*quantity))
            total = total + rowPrice

        table.rows[tableRowCount-1].cells[3].text = str("{:.2f}".format(total))

        for i in range(1,4):
            msDoc.add_paragraph(f"{fd.getRandomItemFromList(sdTextParts)}.")
        fileNameAndPath = os.path.join(outPutPath,documentName)
        msDoc.save(fileNameAndPath)
        #change the create and modify date 
        os.utime(fileNameAndPath, (date.timestamp(), date.timestamp()))

        progressCount = progressCount + 1
        sys.stdout.write(f"({progressCount}/{numberToPrint})  contract progress: %d%%  \r" % (progressCount/numberToPrint*100) )
        sys.stdout.flush()
    print(f"({progressCount}/{numberToPrint})  contract progress: %d%%  \r" % (progressCount/numberToPrint*100) )

def outputApprovalLetter(sequence=0, applicationDate = datetime.datetime.now(), 
                         applicationType="", firstName="",lastName = ""):
    
    date = applicationDate+datetime.timedelta(days=random.choice(range(2,42,1)))
    year =  date.strftime("%Y")
    month = date.strftime("%b")
    batch = date.strftime("")
    documentName = f"LApproval{year}-{month}.{sequence}B{batchName}.docx"
    
    msDoc = Document()
    
    msDoc.add_paragraph(f"{firstName} {lastName}")
    city = fd.getRandomItemFromList(sdCityDictList)
    msDoc.add_paragraph(f"{random.choice(range(2000,3000,10))}, {fd.getRandomItemFromList(sdStreetNames)}")
    msDoc.add_paragraph(f"{city[cCityFields.CityName]}, {city[cCityFields.Province]}")
    msDoc.add_paragraph(f"{city[cCityFields.PostalCode]}")
    msDoc.add_paragraph(f"")
    msDoc.add_paragraph(f"Re: Approval {applicationType} application")
    msDoc.add_paragraph(f"")
    msDoc.add_paragraph(f"")
    tableRowCount = 5
    table = msDoc.add_table(rows=tableRowCount, cols=4)
    #header row
    row_cells = table.rows[0].cells
    row_cells[0].text = "item"
    row_cells[1].text = "quantity"
    row_cells[2].text = "price"
    row_cells[3].text = "" 
    total = 0.00

    for i in range(1,4):
        randomParagraph = ""
        for x in range(2,random.choice(range(3,6))):
            randomParagraph = randomParagraph + f"{fd.getRandomItemFromList(sdTextParts)}. "
        msDoc.add_paragraph(f"{randomParagraph}.")

    paragraph = msDoc.add_paragraph(f"")
    paragraph = msDoc.add_paragraph(f"")
    paragraph = msDoc.add_paragraph(f"Sincerely")
    paragraph = msDoc.add_paragraph(f"")
    paragraph = msDoc.add_paragraph(f"")
    paragraph = msDoc.add_paragraph(f"Admin Clerk {fd.getRandomAtoZ()}")
    paragraph = msDoc.add_paragraph(f"")
    
    fileNameAndPath = os.path.join(outPutPath,documentName)
    msDoc.save(fileNameAndPath)
    #change the create and modify date 
    os.utime(fileNameAndPath, (date.timestamp(), date.timestamp()))

def outputRejectionLetter(sequence=0, applicationDate = datetime.datetime.now(), 
                         applicationType="", firstName="",lastName = ""):
    
    date = applicationDate+datetime.timedelta(days=random.choice(range(2,42,1)))
    year =  date.strftime("%Y")
    month = date.strftime("%b")
    batch = date.strftime("")
    documentName = f"LApproval{year}-{month}.{sequence}B{batchName}.docx"
    
    msDoc = Document()
    
    msDoc.add_paragraph(f"{firstName} {lastName}")
    city = fd.getRandomItemFromList(sdCityDictList)
    msDoc.add_paragraph(f"{random.choice(range(2000,3000,10))}, {fd.getRandomItemFromList(sdStreetNames)}")
    msDoc.add_paragraph(f"{city[cCityFields.CityName]}, {city[cCityFields.Province]}")
    msDoc.add_paragraph(f"{city[cCityFields.PostalCode]}")
    msDoc.add_paragraph(f"")
    msDoc.add_paragraph(f"Re: Rejection {applicationType} application")
    msDoc.add_paragraph(f"")
    msDoc.add_paragraph(f"")
    tableRowCount = 5
    table = msDoc.add_table(rows=tableRowCount, cols=4)
    #header row
    row_cells = table.rows[0].cells
    row_cells[0].text = "item"
    row_cells[1].text = "quantity"
    row_cells[2].text = "price"
    row_cells[3].text = "" 
    total = 0.00

    for i in range(1,8):
        randomParagraph = ""
        for x in range(2,random.choice(range(3,6))):
            randomParagraph = randomParagraph + f"{fd.getRandomItemFromList(sdTextParts)}. "
        msDoc.add_paragraph(f"{randomParagraph}.")

    paragraph = msDoc.add_paragraph(f"")
    paragraph = msDoc.add_paragraph(f"")
    paragraph = msDoc.add_paragraph(f"Sincerely")
    paragraph = msDoc.add_paragraph(f"")
    paragraph = msDoc.add_paragraph(f"")
    paragraph = msDoc.add_paragraph(f"Supervisor {fd.getRandomAtoZ()}")
    paragraph = msDoc.add_paragraph(f"")
    
    fileNameAndPath = os.path.join(outPutPath,documentName)
    msDoc.save(fileNameAndPath)
    #change the create and modify date 
    os.utime(fileNameAndPath, (date.timestamp(), date.timestamp()))

def outputApplications(numberOfApplications = 1):
    progressCount = 0
    for sequence in range(0,numberOfApplications):
        reportElements = []
        date = datetime.datetime.now()+datetime.timedelta(weeks= -1*random.choice(range(1,300,1)))
        year =  date.strftime("%Y")
        month = date.strftime("%b")
        reportName = f"A{year}-{month}.{sequence}B{batchName}.pdf"
        
        applicationType = fd.getRandomItemFromList(sdApplicationTypes)

        report = SimpleDocTemplate(os.path.join(outPutPath,f"{reportName}"))
        styles = getSampleStyleSheet()

        #title
        reportElements.append(Paragraph(f"Application: {applicationType}",styles["h1"]))
        #intro paragraph
        reportElements.append(Paragraph(f"{fd.getRandomItemFromList(sdTextParts)}.{fd.getRandomItemFromList(sdTextParts)}.{fd.getRandomItemFromList(sdTextParts)}.{fd.getRandomItemFromList(sdTextParts)}."))

        firstName = fd.getRandomItemFromList(sdFirstNames)
        lastName = fd.getRandomItemFromList(sdLastNames)
        table_data = [["First Name:",firstName]]
        table_data.append(["Last Name:",lastName])
        dob = datetime.datetime.now()+datetime.timedelta(days= -1*random.choice(range(365*18,365*99)))
        table_data.append(["Date of Birth:",dob.strftime("%Y %b %d")])

        if applicationType in ["Employment",
                        "Daycare Supplement",
                        "Driver's Licence"]:
            makeRandomMoodFace()
            
            reportElements.append(Paragraph(f"Application Image",styles["h2"]))
            #add a applicant image
            attachment_path = "./SeedData/moodface.png"
            attachment_filename = os.path.basename(attachment_path)
            report_image = Image(attachment_path, width=1*inch,height=1*inch)
            report_image.hAlign = "LEFT"
            reportElements.append(report_image)

        if applicationType in ["Dog Licence"]:
            reportElements.append(Paragraph(f"Dog's Picture",styles["h2"]))
            #add a applicant image
            attachment_path = random.choice(sdListOfDogPictures)
            attachment_filename = os.path.basename(attachment_path)
            report_image = Image(attachment_path, width=1*inch,height=1*inch)
            report_image.hAlign = "LEFT"
            reportElements.append(report_image)

        report_table = Table (data=table_data)

        reportElements.append(Paragraph(f"Application Details",styles["h2"]))
        reportElements.append(report_table)

        if applicationType not in ["Housing",
                                "Learning Grant",
                                "Employment",
                                "Daycare Supplement",
                                "Rental Assistance"]:
            expiryDate = datetime.datetime.now()+datetime.timedelta(days= random.choice(range(3,365*2)))
            reportElements.append(Paragraph("Fee Collection",styles["h2"]))
            reportElements.append(Paragraph(f"____________________________________________"))
            reportElements.append(Paragraph(f"Credit card: {fd.getRandomItemFromList(sdCreditCards)}"))
            reportElements.append(Paragraph(f"CVC: 12{fd.getRandom0to9Str}"))
            reportElements.append(Paragraph(f"Expiry date: {expiryDate.strftime('%Y %b %d')}"))
            reportElements.append(Paragraph(f"Amount: {str(random.choice(range(100,900,10)))}"))  
            reportElements.append(Paragraph(f"____________________________________________")) 
        else:
            reportElements.append(Paragraph("Conditions",styles["h2"])) 
            #build approval letter
        
        #report conclusion
        reportElements.append(Paragraph(f""))
        reportElements.append(Paragraph(f""))
        reportElements.append(Paragraph(f"{fd.getRandomItemFromList(sdTextParts)}.{fd.getRandomItemFromList(sdTextParts)}.{fd.getRandomItemFromList(sdTextParts)}.{fd.getRandomItemFromList(sdTextParts)}."))
        
        report.build(reportElements)
        os.utime(os.path.join(outPutPath,f"{reportName}"), (date.timestamp(), date.timestamp()))

        if random.choice(["Reject","Approve","Approve","Approve"]) == "Approve":
            outputApprovalLetter(sequence=sequence,applicationDate=date, 
                                        applicationType=applicationType,firstName=firstName, lastName=lastName)
        else:
            outputRejectionLetter(sequence=sequence,applicationDate=date, 
                                        applicationType=applicationType,firstName=firstName, lastName=lastName)
        
        progressCount = progressCount + 1
        sys.stdout.write(f"({progressCount}/{numberToPrint})  application progress: %d%%  \r" % (progressCount/numberToPrint*100) )
        sys.stdout.flush()   
    sys.stdout.write(f"({progressCount}/{numberToPrint})  application progress: %d%%  \r" % (progressCount/numberToPrint*100) )

def outputSiteVisitImages(numberToPrint = 1):
    progressCount = 0
    for sequence in range (0,numberToPrint):
        date = datetime.datetime.now()+datetime.timedelta(weeks= -1*random.choice(range(1,300,1)))
        year =  date.strftime("%Y")
        month = date.strftime("%b")
        batch = date.strftime("")
        documentName = f"img{year}-{month}.{sequence}B{batchName}"
        makeBuildingSitePhoto(documentName=documentName)

        progressCount = progressCount + 1
        sys.stdout.write(f"({progressCount}/{numberToPrint})  site visit image progress: %d%%  \r" % (progressCount/numberToPrint*100) )
        sys.stdout.flush()
    print(f"({progressCount}/{numberToPrint})  site visit image progress: %d%%  \r" % (progressCount/numberToPrint*100) )

def outputReports():
    
    progressCount = 0
    numberToPrint = len(sdApplicationTypes)

    for application in sdApplicationTypes:
        reportName = f"rpt{application}B{batchName}.pdf"
        df = pd.DataFrame(columns = ["Year","Application Count"])
        
        for year in range (2000,datetime.datetime.now().year):
            df.loc[df.shape[0]] = [year,random.choice(range (200, 2000))]
        
        outputPathAndFileName = os.path.join(outPutPath,f"{reportName}.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = "data"
        for r in dataframe_to_rows(df, index=True, header=True):
            ws.append(r)
        
        for cell in ws['A'] + ws[1]:
            cell.style = 'Pandas'
        
        wsPlot = wb.create_sheet("Plot",index=1)

        # create data for plotting
        values = Reference(ws, min_col = 3, min_row = 1,
                                max_col = 3, max_row = datetime.datetime.now().year - 2000)
        
        # Create object of BarChart class
        chart = BarChart()
        
        # adding data to the Bar chart object
        chart.add_data(values)
        
        # adding x-axis
        cats = Reference(ws, min_col = 2, min_row = 1,
                                max_col = 2, max_row = datetime.datetime.now().year - 2000)
        chart.set_categories(cats)
    
        # set the title of the chart
        chart.title = f"{application} by year"
        
        # set the title of the x-axis
        chart.x_axis.title = "Year"
        
        # set the title of the y-axis
        chart.y_axis.title = "Count"
        
        # add chart to the sheet
        # the top-left corner of a chart
        # is anchored to cell E2 .
        wsPlot.add_chart(chart, "E2")

        wb.save(outputPathAndFileName)
        #df.to_excel(outputPathAndFileName,sheet_name=application, index=False)
        
        progressCount = progressCount + 1
        sys.stdout.write(f"({progressCount}/{numberToPrint})  report progress: %d%%  \r" % (progressCount/numberToPrint*100) )
        sys.stdout.flush()
    print(f"({progressCount}/{numberToPrint})  report progress: %d%%  \r" % (progressCount/numberToPrint*100) )


if __name__ == '__main__':
    #delete the exiting content
    
        
    

    for file in os.listdir(outPutPath):
        fileAndPath = os.path.join(outPutPath,file)
        if os.path.isfile(fileAndPath):
            os.remove(fileAndPath)
    try:
        numberToPrint = int(sys.argv[1])
    except:
        numberToPrint = int(input("how many files should be generated? "))

    dateStart = datetime.datetime.now()

    makeRandomMoodFace()
    print("loading seed data")
    loadSeedData()

    outputContracts(numberToPrint)
    print("")
    
    outputApplications(numberToPrint)
    print("")
    
    outputSiteVisitImages(numberToPrint)
    print()

    outputReports()
    print("")
    
    print("execution time ", format((datetime.datetime.now()-dateStart).total_seconds(),".2f"), "seconds.")
    
else:
    print("nothing")
    